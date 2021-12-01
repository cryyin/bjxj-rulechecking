import docx
import pickle
from collections import defaultdict
# from win32com import client as wc
import re, os, json
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import openpyxl
import xlsxwriter

def read_regulation(path='./规范/规范、标准.txt'):
    regus=defaultdict(list)
    with open(path, 'r', encoding='utf8') as f:
        f.readline()
        lines=f.readlines()
        for line in lines:
            words=line.split()
            # if words[1] not in regus:
            #     regus[words[1]]=words[2]
            # else:
            regus[words[1]].append(words[2])
    return regus
             
def read_para():
    path = "./磁各庄站主体围护结构计算书.docx"
    file = docx.Document(path)
    for i,p in enumerate(file.paragraphs):
        print(i,p.text)

def read_table():
    path = "./磁各庄站主体围护结构计算书.docx"
    file = docx.Document(path)
    for table in file.tables:
        for row in table.rows:
            print("| ",end="")
            for cell in row.cells:
                print(cell.text.replace("\n",""),end=" | ")
            print("\n")

def read_calculation(path):
    
    try:
        file = docx.Document(path)
    except:
        print("计算书打开失败")
        return 'open failed'




    text_order,num_order = extract_framework(path)
    title2table = extract_title2table(path,text_order,num_order)

    data = {}


    ############################################################################
    # 1.	设计依据
    flags = []
    for i,p in enumerate(file.paragraphs):
        if p.text.endswith("设计依据") or p.text.endswith("计算依据") or p.text.endswith("设计遵照的规范") or p.text.endswith("主要规范、标准") or p.text.endswith("遵循的规范与标准") or p.text.endswith("采用的规范") or p.text.endswith("主要规范"):
            flags.append(i)

    # print(flags)
    data["设计依据"] = []
    for flag in flags:
        for i,cont in enumerate(range(flag+1,flag+30)):
            if "《" in file.paragraphs[cont].text:
                data["设计依据"].append("《"+file.paragraphs[cont].text.split("《")[1])

    ############################################################################
    # 3.	荷载参数
    for i,p in enumerate(file.paragraphs):
        if "周边超载" in p.text or "地面超载" in p.text:
            data["荷载参数"] = p.text
            break

    ############################################################################
    # 4.	使用软件
    for i,p in enumerate(file.paragraphs):
        if "理正深" in p.text or "曙光启明星" in p.text:
            data["使用软件"] = p.text
            break

    ############################################################################
    # 5.	设计标准
    data["设计标准"] = ""
    for i,p in enumerate(file.paragraphs):
        if "地下结构的基坑支护结构" in p.text:
            data["设计标准"] += p.text
            break

    for i,p in enumerate(file.paragraphs):
        if "地下结构应按抗浮设防水位" in p.text:
            data["设计标准"] += p.text
            break

    for i,p in enumerate(file.paragraphs):
        if "砂性土地层的侧向水" in p.text:
            data["设计标准"] += p.text
        if "粘性土地层的侧向水" in p.text:
            data["设计标准"] += p.text

    ############################################################################
    # 6.    钢支撑计算
    data["钢支撑计算"] = {}

    flags = []
    for i,p in enumerate(text_order):
        if "支撑计算" in p or "支撑验算" in p or "钢支撑" in p:
            flags.append(i)
    if flags!= []:
        flags.append(flags[-1]+1)
    print("@@@@@@@@@@@@@@@@@@",flags)
    #for flag in flags:
    #    print(text_order[flag])

    # 标题就是钢支撑相关的
    if len(flags)>1:
        ttypes = ["标准段","加宽段","斜撑段"]
        check_types = ["钢支撑刚度验算","钢支撑强度验算","钢支撑整体稳定性验算","钢支撑挠度验算"]
        check_names = ["刚度","强度","稳定性","挠度"]

        target_tt = ""
        for flag in flags:
            for tp in ttypes:
                if flag in text_order and tp in text_order[flag]:
                    target_tt = tp
                    D = []
                    if "剖面" in text_order[flag]:
                        D = re.findall(r"\d-\d剖面",text_order[flag])
                    if D!=[]:
                        target_tt += ("（对应"+D[0]+"）")
                    if target_tt not in data["钢支撑计算"]:
                        data["钢支撑计算"][target_tt] = {}
                    break
            if len(num_order)>flag+1:
                if flag+1 in num_order and flag in num_order and (num_order[flag+1] - num_order[flag])>1:
                    starts = -1
                    ends = -1
                    for op,ct in enumerate(check_types):
                        if ct in text_order[flag]:
                            starts = num_order[flag]
                            ends = num_order[flag+1]
                            ck_name = check_names[op]
                            if target_tt!= "" and ck_name not in data["钢支撑计算"][target_tt]:
                                data["钢支撑计算"][target_tt][ck_name] = ""
                            break
                    if starts!=-1 and ends!=-1:
                        for i,cont in enumerate(range(starts+1,ends)):
                            if target_tt!="":
                                data["钢支撑计算"][target_tt][ck_name] += file.paragraphs[cont].text

    if data["钢支撑计算"]=={}:
        for title in title2table:
            if "剖面" in title:
                Dk = re.findall(r"\d-\d剖面",title)
            for table in title2table[title]:
                for i,row in enumerate(table.rows):
                    for j,cell in enumerate(row.cells):
                        if "强度验算" in cell.text:
                            #print(title,i,j,cell.text,table.rows[i+7].cells [j+2].text)
                            if 'Dk' in locals() or 'Dk' in globals() and Dk!=[]:
                                for part in Dk:
                                    if part not in data["钢支撑计算"]:
                                        data["钢支撑计算"][part] = {}
                                    data["钢支撑计算"][part]["刚度"] = ("刚度验算："+table.rows[i+7].cells [j+2].text)
                        if "平面内稳定性验算" in cell.text:
                            #print(title,i,j,cell.text,table.rows[i+11].cells [j+2].text)
                            if 'Dk' in locals() or 'Dk' in globals() and Dk!=[]:
                                for part in Dk:
                                    if part not in data["钢支撑计算"]:
                                        data["钢支撑计算"][part] = {}
                                    if "稳定性" not in data["钢支撑计算"][part]:
                                        data["钢支撑计算"][part]["稳定性"] = ""
                                    data["钢支撑计算"][part]["稳定性"] += ("平面内稳定性验算：" + table.rows[i+11].cells [j+2].text+"；")
                        if "平面外稳定性验算" in cell.text:
                            #print(title,i,j,cell.text,table.rows[i+9].cells [j+2].text)
                            if 'Dk' in locals() or 'Dk' in globals() and Dk!=[]:
                                for part in Dk:
                                    if part not in data["钢支撑计算"]:
                                        data["钢支撑计算"][part] = {}
                                    if "稳定性" not in data["钢支撑计算"][part]:
                                        data["钢支撑计算"][part]["稳定性"] = ""
                                    data["钢支撑计算"][part]["稳定性"] += ("平面外稳定性验算：" + table.rows[i+9].cells [j+2].text+"；")

    ############################################################################
    # 8.    钢腰梁及连系梁计算
    data["钢腰梁及连系梁计算"] = {}
    key_words = ["钢腰梁","冠梁","钢围檩"]

    flags = []
    for i,p in enumerate(text_order):
        for kw in key_words:
            if kw in p:
                flags.append(i)
    #print(flags)
    starts = -1
    ends = -1
    for flag in flags:
        #print("START:",num_order[flag],file.paragraphs[num_order[flag]].text)
        #print("END:",num_order[flag+1],file.paragraphs[num_order[flag+1]].text)
        if len(num_order) > flag+1:
            if (num_order[flag+1] - num_order[flag])>1:
                starts = num_order[flag]
                ends = num_order[flag+1]
        else:
            starts = num_order[flag]
            ends = len(file.paragraphs)

        if starts!=-1 and ends!=-1:
            data["钢腰梁及连系梁计算"] = {}
            flagt1 = -1
            flagt2 = -1
            for i,cont in enumerate(range(starts,ends)):
                if "抗弯" in file.paragraphs[cont].text:
                    flagt1 = starts+i
                if "抗剪" in file.paragraphs[cont].text:
                    flagt2 = starts+i
            
            if flagt1 != -1 and flagt2 != -1:
                data["钢腰梁及连系梁计算"]["抗弯"] = ""
                for i,cont in enumerate(range(flagt1,flagt2)):
                    data["钢腰梁及连系梁计算"]["抗弯"] += file.paragraphs[cont].text
                data["钢腰梁及连系梁计算"]["抗剪"] = ""
                for i,cont in enumerate(range(flagt2,ends)):
                    data["钢腰梁及连系梁计算"]["抗剪"] += file.paragraphs[cont].text
                #print(target_tt,data["8.钢腰梁及连系梁计算"][target_tt],"\n\n")

    ############################################################################
    # 11.   工程材料信息
    flag = -1
    for i,p in enumerate(file.paragraphs):
        if "工程材料" in p.text:
            flag = i
            break

    #data["11.工程材料信息"] = []
    data["钢筋等级"] = []
    data["型号钢"] = []
    if flag != -1:
        for i,cont in enumerate(range(flag+1,flag+30)):
            if not file.paragraphs[cont].text.startswith("（"):
                break
            if "钢筋：" in file.paragraphs[cont].text:
                data["钢筋等级"].append(file.paragraphs[cont].text.split("：")[1])
            if "钢支撑：" in file.paragraphs[cont].text:
                data["型号钢"].append(file.paragraphs[cont].text.split("：")[1])

    # 注意表格里也有
    gangjin_levels = ["HPB300","HRB400","HRB335"]
    gangjin_xinghao = ["Q235","Q345"]

    if data["钢筋等级"] == []:
        for i,p in enumerate(file.paragraphs):
            for gangjin in gangjin_levels:
                if gangjin in p.text and gangjin not in data["钢筋等级"]:
                    data["钢筋等级"].append(gangjin)
        for title in title2table:
            for table in title2table[title]:
                for i,row in enumerate(table.rows):
                    for j,cell in enumerate(row.cells):
                        for gangjin in gangjin_levels:
                            if gangjin in cell.text and gangjin not in data["钢筋等级"]:
                                data["钢筋等级"].append(gangjin)
    
    if data["型号钢"]  == []:
        for i,p in enumerate(file.paragraphs):
            for xinghao in gangjin_xinghao:
                if xinghao in p.text and xinghao not in data["型号钢"]:
                    data["型号钢"].append(xinghao)
        for title in title2table:
            for table in title2table[title]:
                for i,row in enumerate(table.rows):
                    for j,cell in enumerate(row.cells):
                        for xinghao in gangjin_xinghao:
                            if xinghao in cell.text and xinghao not in data["型号钢"]:
                                data["型号钢"].append(xinghao)

    ############################################################################
    # 12.   抗浮计算
    data["抗浮计算"] = ''
    for i,p in enumerate(file.paragraphs):
        if "抗浮安全系数" in p.text:
            if p.text.endswith("："):
                data["抗浮计算"]+=p.text+"\n"+file.paragraphs[i+1].text
            else:
                data["抗浮计算"]+=p.text

    ############################################################################
    # 嵌固深度
    data["嵌固深度"] = {}
    for title in title2table:
        #print(title)
        for table in title2table[title]:
            for i,row in enumerate(table.rows):
                for j,cell in enumerate(row.cells):
                    if "嵌固深度(m)" in cell.text:
                        #print(i,j,cell.text,table.rows[i].cells[j].text,table.rows[i].cells[j+1].text)
                        if "剖面" in title:
                            D = re.findall(r"\d-\d剖面",title)
                            if D!=[]:
                                for part in D:
                                    data["嵌固深度"][part] = table.rows[i].cells[j+1].text
                        else:
                            data["嵌固深度"][title] = table.rows[i].cells[j+1].text
    #print(data["嵌固深度"])
    ############################################################################
    # 钢支撑轴力值
    data["钢支撑轴力"] = {}
    duan_names = ["标准段","斜撑段","加宽段"]

    # 磁各庄站，直接从文本中提取
    for i,p in enumerate(file.paragraphs):
        if "钢支撑轴力" in p.text:
            #print(i,p.text)
            parts = p.text.split("，")
            for part in parts:
                keys = re.findall(r"(第.*道)",part)
                values = re.findall(r"\d+.\d+KN",part)
                title = find_title(text_order,num_order,i)
                new_title = title
                for duan in duan_names:
                    if duan in title:
                        new_title = duan
                        break
                if "剖面" in title:
                    D = re.findall(r"\d-\d剖面",title)
                    if D!=[]:
                        new_title += ("（对应"+D[0]+"）")
                if new_title not in data["钢支撑轴力"]:
                    data["钢支撑轴力"][new_title] = {}
                if keys!=[] and values!=[]:
                    data["钢支撑轴力"][new_title][keys[0]] = values[0]

    # 金安桥站，从表格中提取
    if data["钢支撑轴力"] == {}:
        for title in title2table:
            found_table = 0
            for table in title2table[title]:
                for i,row in enumerate(table.rows):
                    for j,cell in enumerate(row.cells):
                        if "轴力计算标准值" in cell.text:
                            found_table = 1
                            n = j
                            break
                if found_table == 1:
                    for i,row in enumerate(table.rows):
                        for j,cell in enumerate(row.cells):   
                            if "支撑道数" in cell.text:
                                m = j
                                break
                    new_title = title
                    for duan in duan_names:
                        if duan in title:
                            new_title = duan
                            break
                    if "剖面" in title:
                        D = re.findall(r"\d-\d剖面",title)
                        if D!=[]:
                            new_title += ("（对应"+D[0]+"）")
                    if new_title not in data["钢支撑轴力"]:
                        data["钢支撑轴力"][new_title] = {}
                    for i,row in enumerate(table.rows):
                        if i!=0:
                            data["钢支撑轴力"][new_title][table.rows[i].cells[m].text] = int(float(table.rows[i].cells[n].text)*1.414)
    
    # 积水潭站，需要从文本中判断对应的类型
    if data["钢支撑轴力"] == {}:
        type2id = {}
        for i,p in enumerate(file.paragraphs):
            if "第一道" in p.text:
                #print(p.text)
                title = find_title(text_order,num_order,i)
                #print(title,p.text)
                #title = re.findall(r"\d-\d剖面",title)[0]
                title_re = re.findall(r"\d-\d剖面",title)
                #print(title_re)
                if title_re!=[]:
                    title = title_re[0]
                    if title not in type2id:
                        type2id[title] = {}
                    D = re.findall(r"第一道为(D\d+).*?钢支撑",p.text)
                    if D!=[]:
                        if D[0] not in type2id[title]:
                            type2id[title][D[0]] = []
                        type2id[title][D[0]].append("1")
                    D = re.findall(r"第二.*?道为(D\d+).*?钢支撑",p.text)
                    if D!=[]:
                        if D[0] not in type2id[title]:
                            type2id[title][D[0]] = []
                        type2id[title][D[0]].append("2")
                    D = re.findall(r"第.*?三.*?道为(D\d+).*?钢支撑",p.text)
                    if D!=[]:
                        if D[0] not in type2id[title]:
                            type2id[title][D[0]] = []
                        type2id[title][D[0]].append("3")
                    D = re.findall(r"第.*?四.*?道为(D\d+).*?钢支撑",p.text)
                    if D!=[]:
                        if D[0] not in type2id[title]:
                            type2id[title][D[0]] = []
                        type2id[title][D[0]].append("4")
        #print("\n",type2id)

        for i,p in enumerate(file.paragraphs):
            if "最大轴力标准值" in p.text:
                title = find_title(text_order,num_order,i)
                # title = re.findall(r"\d-\d剖面",title)[0]
                title_re = re.findall(r"\d-\d剖面",title)
                if title_re!=[]:
                    title = title_re[0]
                    if title not in data["钢支撑轴力"]:
                        data["钢支撑轴力"][title] = {}
                    D = re.findall(r"(D\d+).*?钢支撑最大轴力标准值为(.*?)(k|K)(n|N)",p.text)
                    #print(title,type2id[title],p.text,D[0],D[0][0],D[0][1])
                    if D!=[]:
                        if D[0][0] in type2id[title]:
                            for order in type2id[title][D[0][0]]:
                                data["钢支撑轴力"][title][order] = int(float(D[0][1])*1.414)
                                #print(title,D[0][0],order,D[0][1])
        #print(data["钢支撑轴力"][title])
    ############################################################################
    # 挡墙的段落
    data["挡土墙"] = {}
    key_words = ["挡土墙"]

    flags = []
    for i,p in enumerate(text_order):
        for kw in key_words:
            if kw in p:
                flags.append(i)
    #print(flags)
    starts = -1
    ends = -1
    for flag in flags:
        #print("START:",num_order[flag],file.paragraphs[num_order[flag]].text)
        #print("END:",num_order[flag+1],file.paragraphs[num_order[flag+1]].text)
        if len(num_order) > flag+1:
            if (num_order[flag+1] - num_order[flag])>1:
                starts = num_order[flag]
                ends = num_order[flag+1]
        else:
            starts = num_order[flag]
            ends = len(file.paragraphs)

        if starts!=-1 and ends!=-1:
            data["挡土墙"] = []
            for i,cont in enumerate(range(starts,ends)):
                data["挡土墙"].append(file.paragraphs[cont].text) 
    #print(data["挡土墙"])    
    ##############################################################################

    data["支护间距"] = {}
    for title in title2table:
        for table in title2table[title]:
            flag = 0
            for row in table.rows:
                for cell in row.cells:
                    if "支锚类型" in cell.text:
                        flag = 1
                        break
            if flag == 1:
                col_key=0
                col_value=0
                distance_dict={}
                for i, row in enumerate(table.rows):
                    if i==0:
                        for j, cell in enumerate(row.cells):
                            if "道号" in cell.text:
                                col_key=j

                            if '水平间距' in cell.text:
                                col_value=j 
                    else:
                        if re.match(r'\d', row.cells[col_key].text):
                            distance_dict[row.cells[col_key].text]=row.cells[col_value].text
                if "剖面" in title:
                    D = re.findall(r"\d-\d剖面",title)
                    if D!=[]:
                        data["支护间距"][D[0]]=distance_dict
                else:
                    data["支护间距"][title]=distance_dict
    
    print(path,"\n")
    for kk in data.keys():
        #if "设计依据" in kk:
        print(kk," @@ ",data[kk],"\n")
    print("############################################################################")
    return data

# def doc2docx(path):

#     w = wc.Dispatch('Word.Application')
#     # w = wc.DispatchEx('Word.Application')# 或者使用下面的方法，使用启动独立的进程

#     cur_path=os.getcwd()
#     doc=w.Documents.Open(path)
#     doc.SaveAs(path[:-3]+'docx',16) #必须有参数16，否则会出错.

def doc2docx(path):
    out_path = path.replace(path.split("/")[-1],"")
    # print(out_path)
    os.system('libreoffice --headless --convert-to docx --outdir '+out_path+' '+path)

def read_CAD_results(path):
    # 读取CAD上的识图结果
    # path: path to json file

    contents={} # dictionary

    for dir0 in os.listdir(path):
        for dir in os.listdir(os.path.join(path, dir0)):
            if dir=='data' or dir=='text':
                dir_dict={}
                for file in os.listdir(os.path.join(path, dir0, dir)):
                    with open(os.path.join(path, dir0, dir, file), 'r', encoding='utf8') as f:
                        content=json.load(f)
                        dir_dict[file.split('.')[0]]=content
            
                contents[dir0]=dir_dict
            elif dir=='table':
                dir_dict={}
                for file in os.listdir(os.path.join(path, dir0, dir)):
                    with open(os.path.join(path, dir0, dir, file), 'r', encoding='utf8') as f:
                        content=json.load(f)
                        dir_dict[file.split('.')[0]]=content
                
                if 'table' not in contents:
                    contents['table']=dir_dict
                else:
                    contents['table'].update(dir_dict)
    
    return contents
    
def extract_framework(path):
    
    #path = "./doc2/"+filename
    file = docx.Document(path)

    data = {}

    ############################################################################
    text_order = []
    num_order = []

    starts = ["一","二","三","四","五","六","七","八","九","十"]
    for i,p in enumerate(file.paragraphs):
        #print(i,p.text,p.style.name)
        D = re.findall(r"^\d+[^)、）．-][\.?\d+]*[\.?\d]*",p.text)
        if D!=[] and "号线" not in p.text and "年" not in p.text and len(p.text.split(" "))<4:
            # print(i,p.text)
            if i not in num_order:
                text_order.append(p.text)
                num_order.append(i)
        
        for tt in starts:
            if p.text.startswith(tt):
                # print(i,p.text)
                if i not in num_order:
                    text_order.append(p.text)
                    num_order.append(i)
        
        if p.style.name == 'Heading 1':
            # print(i,'Heading 1',p.text)
            if i not in num_order:
                text_order.append(p.text)
                num_order.append(i)
        if p.style.name == 'Heading 2':
            # print(i,'Heading 2',p.text)
            if i not in num_order:
                text_order.append(p.text)
                num_order.append(i)
        if p.style.name == 'Heading 3':
            # print(i,'Heading 3',p.text)
            if i not in num_order:
                text_order.append(p.text)
                num_order.append(i)
    #print(text_order)
    #print(num_order)
    return text_order,num_order

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
 
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

def extract_title2table(path,text_order,num_order):
    title2table = {}
    doc = docx.Document(path)
    title = path.replace(".docx","")
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text in text_order and block.text!="计算方法":
                #print(block.text)
                title = block.text
        elif isinstance(block, Table):
            #print(block)
            if title not in title2table:
                title2table[title] = []
            title2table[title].append(block)
    #print(title2table.keys())
    return title2table

def find_title(text_order,num_order,i):
    for kk,no in enumerate(num_order):
        if i<no and kk!=0:
            return text_order[kk-1]



if __name__ == '__main__':


    files = os.listdir("图纸v1")
    for file in files:
        print(file)
        path_current=os.getcwd()
        path_target=os.path.join(path_current, "图纸v1/"+file)

        ####################
        # 读取规范
        path_regu="./规范/规范、标准.txt"
        regus=read_regulation(path_regu)

        ####################
        # 读取计算书
        path_calc=os.path.join(path_target, "计算书")
        formats={}  # used to indicate whether the file is found, first find .docx, if not find .doc
        for filename in os.listdir(path_calc):
            if not re.match("^~", filename):
                if filename.endswith("docx"):
                    formats['docx']=filename
                elif filename.endswith("doc"):
                    formats['doc']=filename
        
        if 'docx' in formats:
            filename_calc=os.path.join(path_calc, formats['docx'])
        elif 'doc' in formats:
            filename_calc=os.path.join(path_calc, formats['doc'])
        else:
            raise(path_calc+"中没有找到计算书！")
        
        # 判断计算书的格式，如果是.doc结尾的话转为.docx格式
        if filename_calc.endswith(".doc"):
            doc2docx(filename_calc)
            filename_calc=filename_calc[:-3]+"docx"

        print(filename_calc)
        data_calc=read_calculation(filename_calc)

